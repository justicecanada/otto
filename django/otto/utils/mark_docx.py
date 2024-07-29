"""
Recursively search for .docx files in the given directory and add headers/footers to them.
By default, the script will replace the header and footer in document with the text "AI-generated content".
Arguments:
    dir: The directory to search for .docx files. (positional, required)
    mark_text: The text to add to the header and footer. Default is "AI-generated content". (2nd position, optional)
    areas: The areas to add the mark text to. Default is ["header", "footer"]. (3rd position, optional)
"""

import argparse
import os

from docx import Document
from docx.shared import Pt, RGBColor


def mark_docx(docx_path, mark_text="AI-generated content", areas=["header", "footer"]):
    doc = Document(docx_path)
    for section in doc.sections:
        for el in [getattr(section, area) for area in areas]:
            if el is not None:
                el.is_linked_to_previous = False
                el.paragraphs[0].clear()
                run = el.paragraphs[0].add_run(mark_text)
                run.bold = True
                run.font.color.rgb = RGBColor(128, 128, 128)  # Set RGB color to grey
                run.font.size = Pt(18)  # Set font size to 24
                el.paragraphs[0].alignment = 1
    doc.save(docx_path)


def main():
    parser = argparse.ArgumentParser(
        description="Recursively search for .docx files in the given directory and add headers/footers to them."
    )
    parser.add_argument("dir", help="The directory to search for .docx files.")
    parser.add_argument(
        "mark_text",
        nargs="?",
        default="AI-generated content",
        help="The text to add to the header and/or footer. Default is 'AI-generated content'.",
    )
    parser.add_argument(
        "areas",
        nargs="*",
        default=["header", "footer"],
        help="The areas to add the mark text to. Default is ['header', 'footer'].",
    )
    args = parser.parse_args()
    for root, dirs, files in os.walk(args.dir):
        for file in files:
            if file.endswith(".docx"):
                mark_docx(os.path.join(root, file), args.mark_text, args.areas)
