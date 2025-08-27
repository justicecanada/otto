import io
from datetime import datetime

from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from django.urls import reverse
from django.utils import timezone
from django.utils.translation import gettext as _
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_POST

import markdown
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from htmldocx import HtmlToDocx
from rules.contrib.views import objectgetter

from chat.models import Chat
from otto.utils.common import display_cad_cost
from otto.utils.decorators import permission_required


@csrf_exempt
@permission_required("chat.access_chat", objectgetter(Chat, "chat_id"))
def download_chat(request, chat_id):
    chat = get_object_or_404(Chat, id=chat_id)
    title = request.POST.get("title", chat.title)

    # Get the current date
    current_date = timezone.now().strftime("%Y-%m-%d %H:%M:%S")

    # Build the chat URL
    chat_url = request.build_absolute_uri(
        reverse("chat:chat", kwargs={"chat_id": chat_id})
    )

    # Build Otto title
    otto_title = f"Otto - {title}"

    # Create a new Word document
    doc = Document()
    # Set default font to Aptos for all styles and headings to black
    from docx.oxml.ns import qn
    from docx.shared import Pt

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Aptos"
    font.size = Pt(12)

    # Set all heading styles to black and Aptos
    from docx.shared import RGBColor

    for i in range(1, 7):
        style_name = f"Heading {i}"
        if style_name in doc.styles:
            heading_style = doc.styles[style_name]
            heading_style.font.name = "Aptos"
            heading_style.font.color.rgb = RGBColor(0, 0, 0)

    parser = HtmlToDocx()

    # Add document title (Otto heading) in black
    title_para = doc.add_heading(otto_title, 0)
    title_para.runs[0].font.color.rgb = RGBColor(0, 0, 0)

    # Add metadata
    doc.add_paragraph()  # Space

    url_para = doc.add_paragraph()
    url_run = url_para.add_run(f"{_('Chat URL')}: ")
    url_run.bold = True
    url_para.add_run(chat_url)

    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f"{_('Downloaded on')}: ")
    date_run.bold = True
    date_para.add_run(current_date)

    # Add separator
    doc.add_paragraph()
    separator = doc.add_paragraph()
    separator.add_run("─" * 60)
    doc.add_paragraph()

    # Get messages from database
    messages = chat.messages.all().order_by("date_created")

    for message in messages:
        # Add author/date header
        if message.is_bot:
            author = "Otto"
            if message.bot_name:
                author += f" ({message.bot_name})"
            header_color = RGBColor(0, 116, 217)  # Blue
        else:
            author = message.chat.user.full_name if message.chat.user else "User"
            header_color = RGBColor(0, 128, 0)  # Green

        header_text = f"{author} | {message.date_created.strftime('%Y-%m-%d %H:%M:%S')}"

        header_para = doc.add_paragraph()
        header_run = header_para.add_run(header_text)
        header_run.bold = True
        header_run.font.size = Pt(12)
        header_run.font.color.rgb = header_color

        # Add message content
        if message.text and message.text.strip():
            # Choose color for border
            if message.is_bot:
                border_color = "#0074D9"  # Blue
            else:
                border_color = "#008000"  # Green
            # Create HTML with colored border and message text inline
            border_html = (
                f'<span style="color: {border_color}; font-size: 12pt;">▌ </span>'
            )
            content_html = markdown.markdown(message.text)
            # Remove outer <p> tags if present
            if content_html.startswith("<p>") and content_html.endswith("</p>"):
                content_html = content_html[3:-4]
            full_html = f"<span>{border_html}{content_html}</span>"
            # Add to document (will be in the same line)
            parser.add_html_to_document(full_html, doc)

        # Add cost information if available
        if message.usd_cost:
            cost_para = doc.add_paragraph()
            cost_para.paragraph_format.left_indent = Inches(0.5)
            cost_run = cost_para.add_run(f"Cost: {display_cad_cost(message.usd_cost)}")
            cost_run.italic = True
            cost_run.font.size = Pt(9)
            cost_run.font.color.rgb = RGBColor(128, 128, 128)  # Gray

        # Add spacing between messages
        doc.add_paragraph()

    # Save document to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = f"attachment; filename=Otto - {title}.docx"
    return response
