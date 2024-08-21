import io
import json
import mimetypes
import os
import re
import uuid
import zipfile
from datetime import datetime as dt
from io import BytesIO

from django.conf import settings
from django.core.files.base import ContentFile
from django.core.files.storage import default_storage
from django.http import FileResponse, JsonResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.urls import reverse
from django.utils import timezone
from django.utils.translation import gettext as _
from django.views.decorators.http import require_http_methods

import pytesseract
from docx import Document as DocxDocument
from docxtpl import DocxTemplate
from pdf2image import convert_from_path
from structlog import get_logger

from otto.secure_models import AccessKey
from otto.utils.decorators import app_access_required

from .models import Document, Session

logger = get_logger(__name__)

app_name = "case_prep"


@app_access_required(app_name)
def index(request):
    sessions = Session.objects.all(AccessKey(request.user))
    return render(request, "case_prep/index.html", {"sessions": sessions})


@app_access_required(app_name)
def create_session(request):
    session = Session()
    access_key = AccessKey(request.user)
    session.id = uuid.uuid4()
    session.created_by = request.user
    session.name = timezone.now().strftime("%Y-%m-%d %H:%M:%S")
    session.save(AccessKey(bypass=True))
    session.grant_ownership_to(
        access_key,
        modified_by=access_key.user,
        reason="Owner of the object.",
    )
    return redirect("case_prep:session_detail", session_id=session.id)


def session_detail(request, session_id):
    access_key = AccessKey(request.user)
    session = Session.objects.get(access_key, pk=session_id)
    documents = Document.objects.filter(access_key, session=session)
    if documents:
        documents = documents.order_by("sequence")

    return render(
        request,
        "case_prep/session_detail.html",
        {"session": session, "documents": documents},
    )


def upload_files(request):
    # Check if the request method is POST
    if request.method == "POST":

        access_key = AccessKey(request.user)

        # Get the session from the session_id POST variable
        session_id = request.POST.get("session_id")
        session = Session.objects.get(access_key, pk=session_id)

        # Loop through the documents variable in request.FILES
        for uploaded_file in request.FILES.getlist("documents"):
            # If the uploaded file is a ZIP file
            if uploaded_file.name.endswith(".zip"):
                # Extract files from the ZIP archive
                with zipfile.ZipFile(uploaded_file, "r") as zip_ref:
                    for file_name in zip_ref.namelist():
                        # Extract each file and create a Document object
                        with zip_ref.open(file_name) as file:
                            # Get the original path and name of the file within the ZIP
                            original_name = os.path.basename(file_name)
                            # Name the file with the original name but without the path or extension
                            name = os.path.splitext(original_name)[0]
                            # Create a Django File object from the file content
                            django_file = ContentFile(file.read(), name=original_name)
                            # Determine content type based on file extension
                            content_type, _ = mimetypes.guess_type(original_name)
                            # Create a Document object with extracted file
                            if not name:
                                name = original_name
                            document = Document.objects.create(
                                access_key,
                                session=session,  # Assuming session is defined
                                original_name=original_name,
                                name=name,  # You can modify this as needed
                                file=django_file,
                                content_type=content_type or "application/octet-stream",
                            )
            else:
                # Create a Document object for non-ZIP files
                name = os.path.splitext(uploaded_file.name)[0]
                if not name:
                    name = uploaded_file.name
                document = Document.objects.create(
                    access_key,
                    session=session,  # Assuming session is defined
                    original_name=uploaded_file.name,
                    name=name,  # You can modify this as needed
                    file=uploaded_file,
                    content_type=uploaded_file.content_type,  # You can modify this as needed
                )

        documents = Document.objects.filter(access_key, session=session)
        if documents:
            documents = documents.order_by("sequence")

        # Return a the rendering of document_list.html template
        return render(
            request,
            "case_prep/document_list.html",
            {"documents": documents},
        )

    # If request method is not POST, return an error
    return JsonResponse({"error": "Invalid request method."}, status=400)


def delete_session(request, session_id):
    access_key = AccessKey(request.user)
    session = Session.objects.get(access_key, pk=session_id)

    logger.info(f"Deleting session {session_id}.")

    # Delete associated documents
    for document in session.document_set.all(access_key):
        document.delete(access_key)

    # Delete session
    session.delete(access_key)

    # Return success response
    return JsonResponse(
        {"message": "Session deleted successfully.", "url": reverse("case_prep:index")}
    )


def delete_document(request):
    access_key = AccessKey(request.user)

    # Get document_id from the body of the request
    # document_id = json.loads(request.body)["document_id"]
    # document = Document.objects.get(access_key, pk=document_id)
    documents = get_selected_docs(request)

    # logger.info(f"Deleting document {document_id}.")

    # Delete document
    for document in documents:
        logger.info(f"Deleting document {document.id}.")
        document.delete(access_key)
    # document.delete(access_key)

    # Return success response
    return JsonResponse({"message": "Document deleted successfully."})


def save_changes(request):
    access_key = AccessKey(request.user)

    logger.info("Saving changes.")

    if request.method == "POST":

        # Set data to the body of the request which is a JSON array
        data = json.loads(request.body)

        for item in data:
            doc_id = item["id"]
            document = Document.objects.get(access_key, pk=doc_id)
            document.name = item["name"]
            document.sequence = item["sequence"]
            document.hidden = item["hidden"]

            # Check if the item["date"] is a valid YYYY-MM-DD date
            try:
                document.date = dt.strptime(item["date"], "%Y-%m-%d")
            except Exception as e:
                document.date = None

            logger.info(
                f"Saving changes for document {doc_id}.",
                name=document.name,
                sequence=document.sequence,
                hidden=document.hidden,
                date=document.date,
            )

            document.save(access_key)
        return JsonResponse({"message": "Document names updated successfully."})
    else:
        return JsonResponse({"error": "Invalid request method."}, status=400)


# Generate a book of documents by compressing all documents in a session into a ZIP file, naming them according to their sequence number, name, and date
def generate_book_of_documents(request):

    def sanitize_file_name(file_name):
        # Replace spaces and any non-alphanumeric characters with underscores
        sanitized_name = re.sub(r"[^\w\-_\.]", "_", file_name)
        return sanitized_name

    access_key = AccessKey(request.user)

    if request.method == "POST":
        session_id = json.loads(request.body)["session_id"]
        session = Session.objects.get(access_key, pk=session_id)

        # Create a BytesIO stream to hold the ZIP file in memory
        zip_buffer = BytesIO()

        with zipfile.ZipFile(zip_buffer, "w") as zip_ref:
            for document in session.document_set.filter(
                access_key, hidden=False
            ).order_by("sequence"):
                sequence_with_zero_padding = str(document.sequence).zfill(3)
                file_name = f"{sequence_with_zero_padding}_{document.name}_{document.date}.{document.file.name.split('.')[-1]}"
                file_name = sanitize_file_name(file_name)

                # Read the document content
                document_content = document.file.read()

                # Write the document content to the ZIP file
                zip_ref.writestr(file_name, document_content)

        # Ensure the buffer position is at the start
        zip_buffer.seek(0)

        # Save the ZIP file to the desired location in the storage
        file_name = f"{session.created_by.username}-{session.name}.zip"
        file_name = sanitize_file_name(file_name)
        file_path = default_storage.save(file_name, zip_buffer)

        # Update the session model instance with the file path
        session.book_of_documents = file_path
        session.save(access_key)

        return JsonResponse(
            {
                "message": "Book of documents generated successfully.",
                "url": reverse(
                    "case_prep:download_book_of_documents", args=[session_id]
                ),
            }
        )
    else:
        return JsonResponse({"error": "Invalid request method."}, status=400)


def download_book_of_documents(request, session_id):
    access_key = AccessKey(request.user)

    logger.info(f"Downloading book of documents for session {session_id}.")

    session = Session.objects.get(access_key, pk=session_id)

    # Retrieve the File object from the session
    file_obj = session.book_of_documents

    # Return the FileResponse
    response = FileResponse(file_obj.open(), as_attachment=True)
    response["Content-Disposition"] = f"attachment; filename={file_obj.name}"
    return response


def download_document(request, document_id):

    logger.info(f"Downloading document {document_id}.")

    access_key = AccessKey(request.user)
    document = Document.objects.get(access_key, pk=document_id)

    # Return the FileResponse
    response = FileResponse(document.file.open(), as_attachment=True)
    response["Content-Disposition"] = (
        f"attachment; filename={document.name}"  # was original_name before,which is problematic
    )
    return response


def download_documents(request):
    access_key = AccessKey(request.user)

    # Get selected documents using get_selected_docs function
    documents = get_selected_docs(request)

    # Create a BytesIO stream to hold the ZIP file in memory
    zip_buffer = BytesIO()

    with zipfile.ZipFile(zip_buffer, "w") as zip_ref:
        for document in documents:
            logger.info(f"Downloading document {document.id}.")
            # Read the document content
            document_content = document.file.read()
            # Write the document content to the ZIP file
            zip_ref.writestr(document.name, document_content)

    # Ensure the buffer position is at the start
    zip_buffer.seek(0)

    # Return the FileResponse
    response = FileResponse(zip_buffer, as_attachment=True, filename="documents.zip")
    response["Content-Disposition"] = 'attachment; filename="documents.zip"'
    return response


def toggle_document_visibility(request):
    access_key = AccessKey(request.user)

    if request.method == "POST":
        # data = json.loads(request.body)
        # document_id = data["document_id"]
        # document = Document.objects.get(access_key, pk=document_id)
        # document.hidden = not document.hidden
        # document.save(access_key)

        documents = get_selected_docs(request)

        for document in documents:
            document.hidden = not document.hidden
            document.save(access_key)

        return JsonResponse({"message": "Document visibility toggled successfully."})

    return JsonResponse({"error": "Invalid request method."}, status=400)


def create_table_of_contents(request):
    access_key = AccessKey(request.user)

    logger.info("Creating table of contents.")

    session_id = json.loads(request.body)["session_id"]

    # Retrieve the session object
    session = Session.objects.get(access_key, pk=session_id)

    # Load the existing table_of_contents.docx file from the templates directory
    template_path = os.path.join(
        settings.BASE_DIR,
        f"case_prep/templates/case_prep/table_of_contents.docx",
    )
    doc_template = DocxTemplate(template_path)

    # Get all documents associated with the session and order them by sequence
    documents = session.document_set.filter(access_key, hidden=False).order_by(
        "sequence"
    )

    # Loop through the documents and extract the relevant information
    docs = []
    for document in documents:
        docs.append(
            {
                "sequence": document.sequence,
                "name": document.name,
                "date": document.date.strftime("%Y-%m-%d") if document.date else "",
            }
        )

        # Increment the sequence number for each document
        document.sequence += 1
        document.save(access_key)

    # Render the template with the context
    doc_template.render(context={"documents": docs})

    # Create a BytesIO object to act as a temporary file-like object
    output_buffer = io.BytesIO()
    doc_template.save(output_buffer)

    # Write the contents to a file and save to the file system
    doc = ContentFile(output_buffer.getvalue())
    doc.name = "table_of_contents.docx"

    # Create a new Document object with sequence 0 and link it to the Session
    new_document = Document.objects.create(
        access_key,
        session=session,
        original_name=doc.name,
        name="Table of Contents",
        file=doc,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        date=timezone.now(),  # Set the date to the current time
    )

    # Ensure the file is saved before returning the response
    new_document.sequence = 1
    new_document.save(access_key)

    return JsonResponse({"message": "Table of contents created successfully."})


def upvote_feature(request, feature_handle):

    from otto.models import Feedback

    logger.info(f"Upvoting feature: {feature_handle}")

    if request.method == "POST":
        if feature_handle == "qa_research":
            feedback_message = "I am interested in the upcoming Q&A research feature for the Case Prep Assistant. Please upvote it and notify me when it becomes available."
        elif feature_handle == "translate":
            feedback_message = "I would like to see a translation feature in the Case Prep Assistant. Please upvote it and notify me when it becomes available."
        elif feature_handle == "summarization":
            feedback_message = "I am interested in the document summarization feature for the Case Prep Assistant. Please upvote it and notify me when it becomes available."
        else:
            logger.error(f"Invalid feature handle: {feature_handle}")
            raise ValueError("Invalid feature handle")

        feedback = Feedback(
            feedback_type="feedback",
            app=app_name,
            feedback_message=feedback_message,
            modified_by=request.user,
        )
        feedback.save()

        return JsonResponse({"message": _("Feedback submitted successfully")})
    return JsonResponse({"message": _("Invalid request")}, status=400)


def extract_text_from_file(document):
    if document.original_name.lower().endswith(".pdf"):
        images = convert_from_path(document.file.path)  # or name
        text = ""
        for i, image in enumerate(images):
            page_text = pytesseract.image_to_string(image)
            if not page_text.strip():
                print(f"No text found on page {i}")
            else:
                print(f"Text found on page {i}")  #: {page_text[:100]}..."
            text += page_text
        return text
    elif document.original_name.lower().endswith(".docx"):
        doc = DocxDocument(document.file.path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    else:
        raise ValueError("Unsupported file format. Only PDF and DOCX are supported.")


@require_http_methods(["POST"])
def summarize_feature(request):
    from chat.utils import summarize_long_text, tldr_summary

    try:
        access_key = AccessKey(request.user)

        data = json.loads(request.body)
        document_id = data.get("document_id")
        length = data.get("length")
        target_lang = data.get("target_language")

        print(
            f"Received data: document_id={document_id}, length={length}, target_lang={target_lang},"
        )
        documents = get_selected_docs(request)

        # if not document_id:
        #     return JsonResponse({"error": "Document ID is required."}, status=400)
        if not documents:
            return JsonResponse({"error": "No documents selected."}, status=400)

        summarized_texts = []
        for document in documents:
            try:
                file_path = document.file.path
                if not file_path:
                    return JsonResponse({"error": "File path is empty."}, status=400)
                file_contents = extract_text_from_file(document)
                print(f"File contents: {file_contents}")

                summarized_text = summarize_long_text(
                    file_contents, length, target_lang
                )
                print(summarized_text)
                session = document.session
                new_document = create_doc(
                    document, session, summarized_text, access_key, "summarize"
                )
                if not new_document.file:
                    return JsonResponse(
                        {"error": "Failed to create summarized document file."},
                        status=500,
                    )
                summarized_texts.append(summarized_text)
            except Document.DoesNotExist:
                return JsonResponse(
                    {"error": f"Document {document.id} not found."}, status=404
                )

        return JsonResponse({"summarized_texts": summarized_texts})

    except json.JSONDecodeError:
        return JsonResponse({"error": "Invalid JSON data."}, status=400)
    except Exception as e:
        print(f"Unexpected error: {e}")
        return JsonResponse({"error": "An unexpected error occurred."}, status=500)

        # try:
        #     document = Document.objects.get(access_key, pk=document_id)
        #     print("doc found----", document)
        # except Document.DoesNotExist:
        #     return JsonResponse({"error": "Document not found."}, status=404)

    #     file_path = document.file.path
    #     if not file_path:
    #         return JsonResponse({"error": "File path is empty."}, status=400)
    #     file_contents = extract_text_from_file(document)
    #     print(f"File contents: {file_contents}")

    #     summarized_text = summarize_long_text(file_contents, length, target_lang)
    #     print(summarized_text)
    #     session = document.session
    #     new_document = create_doc(
    #         document, session, summarized_text, access_key, "summarize"
    #     )
    #     if not new_document.file:
    #         return JsonResponse(
    #             {"error": "Failed to create summarized document file."}, status=500
    #         )
    #     return JsonResponse({"summarized_text": summarized_text})

    #     # Render the template with the list of documents
    # except json.JSONDecodeError:
    #     return JsonResponse({"error": "Invalid JSON data."}, status=400)
    # except Document.DoesNotExist:
    #     return JsonResponse({"error": "Document not found."}, status=404)
    # except Exception as e:
    #     print(f"Unexpected error: {e}")
    #     return JsonResponse({"error": "An unexpected error occurred."}, status=500)


@require_http_methods(["POST"])
def translate_feature(request):
    from otto.utils.localization import LocaleTranslator

    try:
        access_key = AccessKey(request.user)
        data = json.loads(request.body)
        document_id = data.get("document_id")
        target_lang = data.get("target_language")

        print(f"Received data: document_id={document_id}, target_lang={target_lang}")

        # if not document_id:
        #     return JsonResponse({"error": "Document ID is required."}, status=400)
        if not target_lang:
            return JsonResponse({"error": "Target language is required."}, status=400)

        documents = get_selected_docs(request)

        if not documents:
            return JsonResponse({"error": "No documents selected."}, status=400)

        translated_texts = []

        for document in documents:
            try:
                file_path = document.file.path
                if not file_path:
                    return JsonResponse({"error": "File path is empty."}, status=400)
                file_contents = extract_text_from_file(document)
                file_contents = sanitize_text(file_contents)
                # print(f"File contents: {file_contents}")

                translator = LocaleTranslator(
                    key=settings.AZURE_COGNITIVE_SERVICE_KEY,
                    region=settings.AZURE_COGNITIVE_SERVICE_REGION,
                    endpoint=settings.AZURE_COGNITIVE_SERVICE_ENDPOINT,
                )
                translated_text = translator.translate_text(file_contents, target_lang)
                print(translated_text)
                session = document.session
                new_document = create_doc(
                    document, session, translated_text, access_key, "translate"
                )
                translated_texts.append(translated_text)
            except Document.DoesNotExist:
                return JsonResponse(
                    {"error": f"Document {document.id} not found."}, status=404
                )

        return JsonResponse({"translated_texts": translated_texts})

    except json.JSONDecodeError:
        return JsonResponse({"error": "Invalid JSON data."}, status=400)
    except Exception as e:
        print(f"Unexpected error: {e}")
        logger.error(f"Unexpected error: {e}", exc_info=True)
        return JsonResponse({"error": "An unexpected error occurred."}, status=500)
    #     try:
    #         document = Document.objects.get(access_key, pk=document_id)
    #         # print("Document found:", document)
    #     except Document.DoesNotExist:
    #         return JsonResponse({"error": "Document not found."}, status=404)

    #     file_path = document.file.path
    #     if not file_path:
    #         return JsonResponse({"error": "File path is empty."}, status=400)
    #     file_contents = extract_text_from_file(document)
    #     file_contents = sanitize_text(file_contents)
    #     # print(f"File contents: {file_contents}")

    #     translator = LocaleTranslator(
    #         key=settings.AZURE_COGNITIVE_SERVICE_KEY,
    #         region=settings.AZURE_COGNITIVE_SERVICE_REGION,
    #         endpoint=settings.AZURE_COGNITIVE_SERVICE_ENDPOINT,
    #     )
    #     translated_text = translator.translate_text(file_contents)
    #     print(translated_text)
    #     session = document.session
    #     new_document = create_doc(
    #         document, session, translated_text, access_key, "translate"
    #     )
    #     return JsonResponse({"translated_text": translated_text})

    # except json.JSONDecodeError:
    #     return JsonResponse({"error": "Invalid JSON data."}, status=400)
    # except Document.DoesNotExist:
    #     return JsonResponse({"error": "Document not found."}, status=404)
    # except Exception as e:
    #     print(f"Unexpected error: {e}")
    #     logger.error(f"Unexpected error: {e}", exc_info=True)
    #     return JsonResponse({"error": "An unexpected error occurred."}, status=500)


def create_doc(document, session, text, access_key, feat):

    sanitized_text = sanitize_text(text)

    doc = DocxDocument()

    if feat == "translate":
        doc.add_heading("Translated Document", 0)
    if feat == "summarize":
        doc.add_heading("Summarized Document", 0)
    doc.add_paragraph(sanitized_text)

    # django_file = ContentFile(text, name=document.original_name)
    # # Determine content type based on file extension
    # content_type, _ = mimetypes.guess_type(document.name)
    # new_document_name = f"{feat}d_{document.name}"

    output_buffer = io.BytesIO()
    doc.save(output_buffer)
    output_buffer.seek(0)
    # Create a ContentFile from the BytesIO object
    django_file = ContentFile(output_buffer.read())
    new_document_name = f"{feat}d_{document.name}"
    if not new_document_name.endswith(".docx"):
        new_document_name += ".docx"
    django_file.name = new_document_name

    new_document = Document.objects.create(
        session=session,
        original_name=document.original_name,
        name=new_document_name,  # f"{feat}d_{document.name}",
        file=django_file,  # document.file,  # doc_file,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        date=timezone.now(),
        access_key=access_key,
    )

    # Ensure the file is saved before returning the response
    new_document.sequence = 1
    new_document.save(access_key)
    session.new_document = new_document.file.path
    session.save(access_key)
    return new_document


def sanitize_text(text):
    sanitized_text = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F\x7F\n]", "", text)
    return sanitized_text


def get_selected_docs(request):
    access_key = AccessKey(request.user)
    document_ids = request.POST.getlist("selected_documents")
    documents = [
        get_object_or_404(Document.objects.all(access_key), pk=document_id)
        for document_id in document_ids
    ]
    return documents
