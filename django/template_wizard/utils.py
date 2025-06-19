# Shared helpers for template wizard
import re
from typing import List, Optional

from django.utils.translation import gettext as _

import markdown
from docx import Document
from jinja2 import Template as JinjaTemplate
from llama_index.core.program import LLMTextCompletionProgram
from pydantic import Field, create_model
from structlog import get_logger

from chat.llm import OttoLLM
from chat.utils import url_to_text
from librarian.utils.process_engine import (
    extract_markdown,
    get_process_engine_from_type,
)

logger = get_logger(__name__)

md = markdown.Markdown(extensions=["tables", "extra", "nl2br"], tab_length=2)


def build_pydantic_model_for_fields(
    fields_qs, parent_field=None, model_name="ExtractionTemplate"
):
    """
    Recursively build a Pydantic model for the given fields.
    """
    type_map = {
        "str": str,
        "float": float,
        "int": int,
        "bool": bool,
        "object": None,  # will be replaced by nested model
    }
    fields = {}
    child_fields = fields_qs.filter(parent_field=parent_field)
    for field in child_fields:
        if field.field_type == "object":
            nested_model = build_pydantic_model_for_fields(
                fields_qs,
                parent_field=field,
                model_name=f"{model_name}__{field.slug}",
            )
            base_type = nested_model
        else:
            base_type = type_map.get(field.field_type, str)
        if field.list:
            base_type = List[base_type]
        field_kwargs = {"description": field.description or None}
        if (
            field.field_type == "str"
            and getattr(field, "string_format", "none") != "none"
        ):
            field_kwargs["format"] = field.string_format
        if field.required:
            fields[field.slug] = (
                base_type,
                Field(..., **field_kwargs),
            )
        else:
            fields[field.slug] = (
                Optional[base_type],
                Field(default=None, **field_kwargs),
            )
    return create_model(model_name, **fields)


def unpack_model_to_dict(model_instance):
    """
    Recursively unpack a Pydantic model instance to a dict, ensuring all nested models and lists are handled.
    """
    if hasattr(model_instance, "dict"):
        return unpack_model_to_dict(model_instance.dict())
    elif isinstance(model_instance, dict):
        return {k: unpack_model_to_dict(v) for k, v in model_instance.items()}
    elif isinstance(model_instance, list):
        return [unpack_model_to_dict(item) for item in model_instance]
    else:
        return model_instance


def extract_source_text(source):
    if source.url:
        try:
            source.text = url_to_text(source.url)
            source.save()
        except Exception as e:
            logger.error(
                "Error processing URL",
                url=source.url,
                source_id=source.id,
                error=str(e),
            )
            raise e
    elif source.saved_file:
        try:
            with source.saved_file.file.open("rb") as file:
                content = file.read()
                content_type = source.saved_file.content_type
                process_engine = get_process_engine_from_type(content_type)
                source.text, _ = extract_markdown(
                    content, process_engine, pdf_method=source.session.pdf_method
                )
                source.save()
        except Exception as e:
            logger.error(
                "Error processing saved file",
                filelname=source.filename,
                source_id=source.id,
                error=str(e),
            )
            raise e
    assert source.text, "Source text extraction failed"


def extract_fields(source):
    """
    Uses LLM to extract fields from the source text and saves to source.extracted_json (TextField)
    """

    template = source.session.template
    if not template:
        logger.error(
            "Missing template for extraction",
            source_id=source.id,
        )
        raise Exception("Missing template for extraction")
    if not source.text:
        extract_source_text(source)
    if not template.fields.exists():
        logger.error(
            "Missing fields for extraction",
            source_id=source.id,
        )
        raise Exception("Missing fields for extraction")
    TemplateModel = build_pydantic_model_for_fields(template.fields.all())
    llm = OttoLLM(deployment="gpt-4.1")
    prompt = (
        "Extract the requested fields from this document, if they exist "
        "(otherwise, the field value should be None):\n\n"
        "<document>\n{document_text}\n</document>"
    ).format(document_text=source.text)
    program = LLMTextCompletionProgram.from_defaults(
        llm=llm.llm,
        output_cls=TemplateModel,
        prompt_template_str=prompt,
        verbose=True,
    )
    try:
        result = program(document_text=source.text)
        result_dict = unpack_model_to_dict(result)
        source.extracted_json = result_dict
        source.save()
    except Exception as e:
        logger.error("Error extracting fields", source_id=source.id, error=str(e))
        source.extracted_json = None
        source.save()
        raise


def _convert_markdown_fields(data):
    """
    Recursively convert all string fields in data from markdown to HTML.
    """
    if isinstance(data, dict):
        return {k: _convert_markdown_fields(v) for k, v in data.items()}
    elif isinstance(data, list):
        return [_convert_markdown_fields(item) for item in data]
    elif isinstance(data, str):
        # Convert markdown to HTML only if the string is not empty
        return md.convert(data) if data.strip() else data
    else:
        return data


def fill_template_from_fields(source):
    """
    Fills the template with the extracted fields, saves to source.template_result (TextField)
    """

    template = source.session.template
    if not template or not source.extracted_json:
        logger.error(
            "Missing template or extracted fields for template filling",
            source_id=source.id,
        )
        return
    try:
        fields_data = source.extracted_json
        # Convert all string fields from markdown to HTML
        fields_data = _convert_markdown_fields(fields_data)
    except Exception as e:
        logger.error(
            "Error loading extracted_json for template filling",
            source_id=source.id,
            error=str(e),
        )
        source.template_result = None
        source.save()
        return
    output = None
    if template.layout_jinja:
        try:
            jinja_template = JinjaTemplate(template.layout_jinja)
            output = jinja_template.render(**fields_data)
        except Exception as e:
            logger.error(
                "Error rendering Jinja template", source_id=source.id, error=str(e)
            )
            output = None
    else:
        output = _(
            "Template layout not defined. Please set a Jinja layout for this template."
        )
    source.template_result = output
    source.save()


def validate_docx_template_fields(docx_file, template) -> dict:
    """
    Validates that the docx file contains all top-level TemplateField slugs for the template.
    Returns a dict with keys: is_valid, missing_fields, invalid_fields, found_fields, required_fields.
    """
    required_slugs = template.top_level_slugs
    found_slugs = set()
    invalid_slugs = set()
    # Extract all text from the docx file
    try:
        if hasattr(docx_file, "open"):
            # Django FileField
            file_obj = docx_file.open("rb")
        else:
            file_obj = docx_file
        doc = Document(file_obj)
        text = " ".join([p.text for p in doc.paragraphs])
        # Also check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += " " + cell.text
        # Find all {{ slug }} patterns
        matches = re.findall(r"{{\s*(\w+)\s*}}", text)
        found_slugs = set(matches)
        # Invalid = found but not required
        invalid_slugs = found_slugs - required_slugs
        missing_slugs = required_slugs - found_slugs
        is_valid = not missing_slugs and not invalid_slugs
        return {
            "is_valid": is_valid,
            "missing_fields": sorted(missing_slugs),
            "invalid_fields": sorted(invalid_slugs),
            "found_fields": sorted(found_slugs),
            "required_fields": sorted(required_slugs),
        }
    except Exception as e:
        return {
            "is_valid": False,
            "missing_fields": list(required_slugs),
            "invalid_fields": [],
            "found_fields": [],
            "required_fields": list(required_slugs),
            "error": str(e),
        }
